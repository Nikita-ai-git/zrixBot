import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from docx import Document
import os

# ----------------------------- CONFIG -----------------------------

MAX_PAGES = 20  # Set a cap to avoid infinite crawling
visited = set()  # Track visited URLs
doc = Document()  # Create a new Word document
doc.add_heading("Website Scraping Report", 0)

# ---------------------------- UTILITIES ---------------------------

def is_valid_internal_url(url, base_domain):
    """Only allow URLs from the same domain (internal links)"""
    parsed = urlparse(url)
    return parsed.netloc == "" or parsed.netloc == base_domain

def clean_and_extract_text(soup):
    """Remove non-content elements and return clean text lines"""
    # Remove non-content sections
    for tag in soup(["script", "style", "header", "footer", "nav", "aside"]):
        tag.decompose()

    # Extract all visible text
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    return [line for line in lines if line]  # Remove empty lines

# ---------------------------- MAIN SCRAPER -------------------------

def scrape_page(url, base_url):
    """Recursively scrape page content and follow internal links"""
    if url in visited or len(visited) >= MAX_PAGES:
        return

    print(f"🔗 Scraping: {url}")
    visited.add(url)

    try:
        response = requests.get(url, timeout=10)
        if response.status_code != 200:
            print(f"⚠️ Skipped: {url} (status {response.status_code})")
            return
    except Exception as e:
        print(f"❌ Error fetching {url}: {e}")
        return

    soup = BeautifulSoup(response.text, "html.parser")
    text_lines = clean_and_extract_text(soup)

    # Add heading and text to Word doc
    doc.add_heading(f"Page: {url}", level=1)
    for line in text_lines:
        doc.add_paragraph(line)

    # Extract and follow all internal links
    base_domain = urlparse(base_url).netloc
    for tag in soup.find_all("a", href=True):
        href = tag['href']
        full_url = urljoin(url, href)  # Resolve relative link
        norm_url = full_url.split("#")[0].rstrip("/")  # Normalize URL

        if norm_url not in visited and is_valid_internal_url(full_url, base_domain):
            scrape_page(norm_url, base_url)

# ---------------------------- RUN SCRIPT ---------------------------

if __name__ == "__main__":
    start_url = "https://www.gfgc.com/"
    scrape_page(start_url, start_url)

    # Save result to .docx
    os.makedirs("output", exist_ok=True)
    output_path = "output/scraped_content.docx"
    doc.save(output_path)
    print(f"\n✅ Finished. Saved to: {output_path}")
